from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\User\Downloads\COV\docs\project\Registro_de_prompts_COV.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6472"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_label_body(doc, label, body):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=4)
    label_run = paragraph.add_run(label + ": ")
    set_font(label_run, size=10.5, color=DARK_BLUE, bold=True)
    body_run = paragraph.add_run(body)
    set_font(body_run, size=10.5, color=INK)
    return paragraph


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, **edges):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = "w:" + edge_name
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def add_prompt_callout(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=7, line_spacing=1.18)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    paragraph.paragraph_format.space_before = Pt(4)
    set_paragraph_shading(paragraph, CALLOUT)
    set_paragraph_border(
        paragraph,
        left={"val": "single", "sz": "18", "space": "8", "color": BLUE},
    )
    run = paragraph.add_run(text)
    set_font(run, size=10.5, color=INK, italic=True)


def add_record(doc, number, title, prompt, objective, result, learning):
    heading = doc.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    set_paragraph_spacing(heading, before=14, after=5, line_spacing=1.0)
    run = heading.add_run(f"{number}. {title}")
    set_font(run, size=13, color=BLUE, bold=True)

    label = doc.add_paragraph()
    label.paragraph_format.keep_with_next = True
    set_paragraph_spacing(label, after=3)
    run = label.add_run("PROMPT TEXTUAL")
    set_font(run, size=8.5, color=MUTED, bold=True)

    add_prompt_callout(doc, prompt)
    add_label_body(doc, "Objetivo", objective)
    add_label_body(doc, "Resultado", result)
    add_label_body(doc, "Aprendizaje", learning)


def add_note_box(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=8, line_spacing=1.18)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(4)
    set_paragraph_shading(p, "F8FAFC")
    set_paragraph_border(
        p,
        top={"val": "single", "sz": "6", "space": "0", "color": "D8E2EE"},
        bottom={"val": "single", "sz": "6", "space": "0", "color": "D8E2EE"},
        left={"val": "single", "sz": "6", "space": "0", "color": "D8E2EE"},
        right={"val": "single", "sz": "6", "space": "0", "color": "D8E2EE"},
    )
    title = p.add_run("Nota de integridad. ")
    set_font(title, size=10, color=DARK_BLUE, bold=True)
    body = p.add_run(text)
    set_font(body, size=10, color=INK)


def add_metadata_rows(doc):
    heading = doc.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    set_paragraph_spacing(heading, before=2, after=5, line_spacing=1.0)
    run = heading.add_run("Datos del documento")
    set_font(run, size=13, color=BLUE, bold=True)
    for label, value in (
        ("Proyecto", "COV App"),
        ("Periodo", "Agosto de 2026"),
        ("Herramienta", "Codex (agente de inteligencia artificial)"),
        ("Propósito", "Documentar instrucciones reales usadas durante el desarrollo."),
    ):
        add_label_body(doc, label, value)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header_p, after=0, line_spacing=1.0)
    h1 = header_p.add_run("COV APP")
    set_font(h1, size=8.5, color=DARK_BLUE, bold=True)
    h2 = header_p.add_run("  |  Evidencia académica de prompting")
    set_font(h2, size=8.5, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_p, after=0, line_spacing=1.0)
    footer_run = footer_p.add_run("COV App | Registro de prompts | Agosto 2026")
    set_font(footer_run, size=8.5, color=MUTED)


def build_document():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=12, after=3, line_spacing=1.0)
    title_run = title.add_run("Registro de prompts")
    set_font(title_run, size=24, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12, line_spacing=1.15)
    subtitle_run = subtitle.add_run(
        "Evidencia de las instrucciones realizadas a un agente de IA durante el desarrollo de COV App"
    )
    set_font(subtitle_run, size=12, color=MUTED)

    intro = doc.add_paragraph()
    set_paragraph_spacing(intro, after=10)
    intro_run = intro.add_run(
        "Este registro conserva los mensajes solicitados durante el proyecto y resume el propósito, el resultado y el aprendizaje obtenido en cada interacción."
    )
    set_font(intro_run, size=10.5, color=INK)

    add_metadata_rows(doc)
    add_note_box(
        doc,
        "Los prompts se transcriben tal como fueron escritos, incluidas sus variaciones ortográficas. No se añadieron solicitudes inventadas.",
    )

    records = [
        (
            "Organización de la carpeta del proyecto",
            "QUIERO QUE ORGANIZES LA CARPETA DE COV_APP",
            "Ordenar el proyecto Flutter sin alterar los cambios que ya estaban en desarrollo.",
            "Se reorganizaron los recursos de diseño y los assets, y se documentó la estructura del proyecto.",
            "Una petición corta puede resolverse mejor cuando se valida primero la estructura existente y se evita modificar archivos funcionales sin necesidad.",
        ),
        (
            "Mejora funcional y entrega APK",
            "veras esto tengo que hacer debe de ser una aplicacion 100% funcional mi entregable final es una apk donde se pueda ver todo lo que se necesita ya debo tener almenos hecho una parte del proyecto mira hasta donde tenemos en el cov_app y mejorale o sigue avanzando debe de esatr 100% funcioanl debes de copilar para yo poder ver como vas",
            "Revisar el estado real de COV App, continuar su desarrollo y preparar una versión instalable para demostrar el avance.",
            "Se revisó el plan académico, se reforzó la base funcional de la aplicación y se dejaron compilaciones APK dentro del proyecto como evidencia de avance.",
            "Es útil especificar el entregable final, el criterio de funcionalidad y la necesidad de ver una compilación para orientar las prioridades técnicas.",
        ),
        (
            "Ejecución de la aplicación en Web",
            "ejecuta en el navegador la app",
            "Iniciar la versión Web de COV App para comprobar que la aplicación se puede visualizar en un navegador.",
            "La aplicación se ejecutó en Google Chrome y se verificó respuesta correcta en la dirección local http://localhost:7357.",
            "Solicitar una plataforma concreta de prueba permite validar la experiencia del usuario además de la compilación del código.",
        ),
        (
            "Generación de evidencia de prompting",
            "quiero que me des un archivo de como te eh pedido los promts",
            "Crear un archivo entregable con el registro de las instrucciones usadas durante el desarrollo asistido por IA.",
            "Se generó este documento Word con las solicitudes reales, su contexto y los aprendizajes asociados.",
            "Registrar los prompts desde el inicio ayuda a construir el compendio de buenas prácticas solicitado en el plan del proyecto.",
        ),
    ]

    for index, (title_text, prompt, objective, result, learning) in enumerate(records, start=1):
        add_record(doc, index, title_text, prompt, objective, result, learning)

    conclusion = doc.add_paragraph()
    conclusion.paragraph_format.keep_with_next = True
    set_paragraph_spacing(conclusion, before=16, after=6, line_spacing=1.0)
    c_run = conclusion.add_run("Buenas prácticas identificadas")
    set_font(c_run, size=16, color=BLUE, bold=True)

    lessons = [
        ("Define el resultado esperado", "por ejemplo, una APK, una ejecución Web o un documento de evidencia."),
        ("Incluye el contexto del proyecto", "como el nombre de la aplicación, el estado actual y el plan de trabajo."),
        ("Indica cómo validar el resultado", "por ejemplo, compilar, analizar, ejecutar pruebas o abrir la aplicación en navegador."),
        ("Registra cada interacción", "para poder explicar qué se pidió, qué se logró y cómo mejorar futuras solicitudes."),
    ]
    for label, body in lessons:
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, after=4)
        r1 = paragraph.add_run(label + ". ")
        set_font(r1, size=10.5, color=DARK_BLUE, bold=True)
        r2 = paragraph.add_run(body)
        set_font(r2, size=10.5, color=INK)

    final_note = doc.add_paragraph()
    set_paragraph_spacing(final_note, before=8, after=0)
    final_note_run = final_note.add_run(
        "Documento preparado como respaldo del proceso de desarrollo asistido por inteligencia artificial de COV App."
    )
    set_font(final_note_run, size=9.5, color=MUTED, italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
